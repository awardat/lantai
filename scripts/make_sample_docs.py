"""生成演示用样例文档（仅开发自测，演示文档由用户提供放入 docs/）。

用法：python make_sample_docs.py [输出目录]   # 默认 scripts/sample_docs/
生成：txt / md / docx / pdf（文字）/ 图片（png，用 PIL 可选）样例。
"""
from __future__ import annotations

import sys
from pathlib import Path

OUT_DEFAULT = Path(__file__).resolve().parent / "sample_docs"

SAMPLES = {
    "兰台简介.txt": "兰台是本地运行的 RAG 知识库演示系统。\n\n"
    "兰台之名取自汉代皇家档案馆。汉代兰台令史典校秘书、掌图籍秘书之事，"
    "班固曾为兰台令史，后世遂以兰台代指国家藏书与档案之府。\n\n"
    "本系统支持上传 txt、md、pdf、docx 与图片，自动解析、切块、向量化后入库，"
    "可用自然语言提问并获得带引用来源的答案。\n\n"
    "系统支持 Ollama 本地模型与 OpenAI 兼容云端 API 两种 Provider，"
    "并且可以为不同类型的文件配置不同的 AI 模型。",
    "知识库问答说明.md": "# 知识库问答\n\n"
    "上传文档后，在「问答」页输入问题，系统会检索最相关的切片并生成答案。\n\n"
    "## 引用来源\n\n"
    "每个答案都会附带引用来源，展示文档名与相似度分数，点击可预览源文件。\n\n"
    "## 按类型配置 AI\n\n"
    "图片走视觉模型、扫描件走 OCR 模型、文字文档走文本模型，均可在设置页配置。",
    "会议纪要样例.docx": None,  # 运行时生成
    "产品手册样例.pdf": None,   # 运行时生成
}


def make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("兰台产品手册（样例）", level=1)
    doc.add_paragraph("本文档为开发自测样例，演示 docx 解析与问答检索。")
    doc.add_heading("功能列表", level=2)
    table = doc.add_table(rows=3, cols=2)
    for i, (k, v) in enumerate([("文档管理", "上传/解析/删除"), ("知识问答", "检索+引用"), ("配置", "按类型 AI")]):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.save(str(path))


def make_pdf(path: Path) -> None:
    """无重依赖生成文字 PDF：用 pypdf 写不可行，退化为纯文本说明 + 尝试 reportlab（可选）。"""
    try:
        from reportlab.pdfgen import canvas  # type: ignore

        c = canvas.Canvas(str(path))
        c.setFont("Helvetica", 14)
        c.drawString(60, 780, "Lantai Sample PDF")
        c.setFont("Helvetica", 11)
        text = (
            "This is a sample text PDF for local RAG testing. "
            "Lantai keeps your documents on this machine and answers questions with citations. "
        ) * 12
        y = 750
        for line in [text[i : i + 90] for i in range(0, len(text), 90)]:
            c.drawString(60, y, line)
            y -= 18
        c.save()
    except ImportError:
        # 无 reportlab 时生成一个可被 pypdf 解析的最小 PDF（手工构造）
        content = (
            "BT /F1 12 Tf 60 780 Td (Lantai Sample PDF - text layer) Tj ET "
            "BT /F1 10 Tf 60 750 Td (This PDF has a text layer for RAG testing.) Tj ET "
        )
        objects = (
            "1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
            "2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
            "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n"
            f"4 0 obj << /Length {len(content)} >> stream\n{content}\nendstream endobj\n"
            "5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
            "xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
            "0000000133 00000 n \n0000000242 00000 n \n0000000351 00000 n \n"
            "trailer << /Size 6 /Root 1 0 R >>\nstartxref\n407\n%%EOF\n"
        )
        path.write_bytes(objects.encode("latin-1"))


def make_png(path: Path) -> None:
    """生成一张纯色示例图片（无 PIL 依赖时用最小 PNG 字节）。"""
    try:
        from PIL import Image  # type: ignore

        img = Image.new("RGB", (400, 120), (138, 61, 43))
        img.save(str(path))
    except ImportError:
        # 1x1 红色 PNG
        import base64

        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
        )
        path.write_bytes(png)


def main() -> None:
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else OUT_DEFAULT
    out.mkdir(parents=True, exist_ok=True)
    for name, content in SAMPLES.items():
        p = out / name
        if p.suffix in (".txt", ".md"):
            p.write_text(content or "", encoding="utf-8")
        elif p.suffix == ".docx":
            make_docx(p)
        elif p.suffix == ".pdf":
            make_pdf(p)
        print(f"生成：{p}")
    make_png(out / "示例图片.png")
    print(f"生成：{out / '示例图片.png'}")
    print(f"完成。样例目录：{out}")


if __name__ == "__main__":
    main()
